from __future__ import annotations

import argparse
import csv
import json
import re
import zipfile
from pathlib import Path

from docx import Document
from pypdf import PdfReader


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def document_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def validate(args: argparse.Namespace) -> dict[str, object]:
    required = [args.source, args.main, args.supplement, args.main_pdf, args.supplement_pdf, args.basin_csv]
    required += [args.figure_dir / f"Figure{i}_{name}.png" for i, name in [(4, "CommonReferenceWorkflow"), (5, "DeltaTotal"), (6, "DeltaCC"), (7, "DeltaHA")]]
    missing = [str(path) for path in required if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing deliverable input: " + "; ".join(missing))

    with zipfile.ZipFile(args.source) as source_zip, zipfile.ZipFile(args.main) as main_zip:
        source_xml = source_zip.read("word/document.xml")
        main_xml = main_zip.read("word/document.xml")
        field_counts = {
            "source_cites": len(re.findall(rb"ADDIN\s+EN\.CITE", source_xml)),
            "main_cites": len(re.findall(rb"ADDIN\s+EN\.CITE", main_xml)),
            "source_reflist": len(re.findall(rb"ADDIN\s+EN\.REFLIST", source_xml)),
            "main_reflist": len(re.findall(rb"ADDIN\s+EN\.REFLIST", main_xml)),
            "source_begins": source_xml.count(b'w:fldCharType="begin"'),
            "main_begins": main_xml.count(b'w:fldCharType="begin"'),
            "source_ends": source_xml.count(b'w:fldCharType="end"'),
            "main_ends": main_xml.count(b'w:fldCharType="end"'),
        }
        if list(field_counts.values()) != [32, 32, 1, 1, 33, 33, 33, 33]:
            raise RuntimeError(f"EndNote field mismatch: {field_counts}")

        custom_parts = sorted(name for name in source_zip.namelist() if name.startswith("customXml/") and not name.endswith("/"))
        if custom_parts != sorted(name for name in main_zip.namelist() if name.startswith("customXml/") and not name.endswith("/")):
            raise RuntimeError("customXml part list changed")
        custom_equal = all(source_zip.read(name) == main_zip.read(name) for name in custom_parts)
        if not custom_equal:
            raise RuntimeError("customXml content changed")

        figure_files = [
            args.figure_dir / "Figure4_CommonReferenceWorkflow.png",
            args.figure_dir / "Figure5_DeltaTotal.png",
            args.figure_dir / "Figure6_DeltaCC.png",
            args.figure_dir / "Figure7_DeltaHA.png",
        ]
        media_equal = all(main_zip.read(f"word/media/image{i}.png") == path.read_bytes() for i, path in enumerate(figure_files, 4))
        if not media_equal:
            raise RuntimeError("Embedded main figures do not match final figure exports")

    main_doc = Document(args.main)
    main_text = document_text(main_doc)
    must_have = [
        "Climate and Aggregate Human-Activity Scenario Contrasts in Flood and Drought Event Frequencies across China (1950–2019): An ISIMIP3a Multi-Model Analysis",
        "common factual SRI reference",
        "1,743 have all seven paired models",
        "pooled DIndex values of 0.634–0.864",
        "near parity at 0.478",
        "Data Availability",
        "Code Availability",
    ]
    must_not_have = [
        "[CITATION NEEDED]",
        "[MEANING CHECK]",
        "available files do not",
        "requires author verification",
        "project-review files",
        "before submission",
        "Codex",
        "correlation analysis",
        "Mean_SCI",
        "32.4%",
        "36.9%",
        "34.9%",
        "1,460 have all seven paired models",
        "pooled DIndex values from 0.645 to 0.840",
        "pooled drought DIndex is 0.491",
        "[repository and persistent identifier",
    ]
    absent_required = [text for text in must_have if text not in main_text]
    if absent_required:
        raise RuntimeError(f"Required main-text phrases are absent: {absent_required}")

    basin_rows = read_csv(args.basin_csv)
    basin_lookup = {(row["Basin_Name"], row["Event_Type"]): row for row in basin_rows}
    table = main_doc.tables[0]
    for row in table.rows[1:]:
        basin = row.cells[0].text.strip()
        flood = basin_lookup[(basin, "Flood")]
        drought = basin_lookup[(basin, "Drought")]
        expected = [
            f"{float(flood['Mean_Delta_CC_mean']):.3f} ± {float(flood['Mean_Delta_CC_std']):.3f}",
            f"{float(flood['Mean_Delta_HA_mean']):.3f} ± {float(flood['Mean_Delta_HA_std']):.3f}",
            f"{float(flood['Pooled_DIndex']):.3f}",
            f"{float(drought['Mean_Delta_CC_mean']):.3f} ± {float(drought['Mean_Delta_CC_std']):.3f}",
            f"{float(drought['Mean_Delta_HA_mean']):.3f} ± {float(drought['Mean_Delta_HA_std']):.3f}",
            f"{float(drought['Pooled_DIndex']):.3f}",
        ]
        observed = [cell.text.strip() for cell in row.cells[2:]]
        if observed != expected:
            raise RuntimeError(f"Table 1 mismatch for {basin}: {observed} != {expected}")

    supplement_doc = Document(args.supplement)
    supplement_text = document_text(supplement_doc)
    combined_text = main_text + "\n" + supplement_text
    forbidden_counts = {text: combined_text.count(text) for text in must_not_have}
    stale_present = [text for text, count in forbidden_counts.items() if count]
    if stale_present:
        raise RuntimeError(f"Forbidden or stale text remains: {stale_present}")
    if len(supplement_doc.tables) != 5 or len(supplement_doc.inline_shapes) != 1:
        raise RuntimeError("Supplement structure mismatch")
    for phrase in ["Table S1.", "Table S5.", "Figure S1.", "Public repository identifiers have not yet been assigned"]:
        if phrase not in supplement_text:
            raise RuntimeError(f"Supplement phrase missing: {phrase}")

    pdf_pages = {}
    pdf_texts = {}
    for label, path, expected_pages in [
        ("main", args.main_pdf, 29),
        ("supplement", args.supplement_pdf, 4),
    ]:
        reader = PdfReader(path)
        counts = [len((page.extract_text() or "").strip()) for page in reader.pages]
        if len(reader.pages) != expected_pages or min(counts) == 0:
            raise RuntimeError(f"PDF validation failed for {label}: pages={len(reader.pages)}, text_counts={counts}")
        pdf_pages[label] = {"pages": len(reader.pages), "minimum_page_text_characters": min(counts)}
        pdf_texts[label] = [page.extract_text() or "" for page in reader.pages]

    main_pages = pdf_texts["main"]
    supplement_pages = pdf_texts["supplement"]
    table_page = next((i + 1 for i, text in enumerate(main_pages) if "Table 1. Common-reference" in text), None)
    if table_page is None or not all(
        re.search(rf"\b{basin}\b", main_pages[table_page - 1])
        for basin in ["CSL", "PSE", "SWRB", "YHH", "YZR"]
    ):
        raise RuntimeError("Table 1 is not complete on one PDF page")
    table_s5_page = next((i + 1 for i, text in enumerate(supplement_pages) if "Table S5. Grid-ensemble" in text), None)
    if table_s5_page is None or not all(
        term in supplement_pages[table_s5_page - 1]
        for term in ["Drought", "Flood", "0.882", "-0.225"]
    ):
        raise RuntimeError("Table S5 is not complete on one PDF page")

    figure_pages = {}
    for number in range(1, 8):
        pages = [i + 1 for i, text in enumerate(main_pages) if re.search(rf"\bFigure\s+{number}\b", text)]
        if not pages:
            raise RuntimeError(f"Figure {number} caption not found in PDF")
        figure_pages[str(number)] = pages

    equation_pages = {}
    for number in range(1, 9):
        pages = [i + 1 for i, text in enumerate(main_pages[:20]) if f"({number})" in text]
        if not pages:
            raise RuntimeError(f"Equation ({number}) not found before References")
        equation_pages[str(number)] = pages

    return {
        "status": "PASS",
        "endnote_fields": field_counts,
        "custom_xml_parts": len(custom_parts),
        "custom_xml_unchanged": custom_equal,
        "embedded_figures_match_exports": media_equal,
        "main_table_rows_checked": len(table.rows) - 1,
        "supplement_tables": len(supplement_doc.tables),
        "supplement_figures": len(supplement_doc.inline_shapes),
        "pdfs": pdf_pages,
        "table_1_page": table_page,
        "table_s5_page": table_s5_page,
        "figure_caption_pages": figure_pages,
        "equation_pages": equation_pages,
        "forbidden_counts": forbidden_counts,
    }


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--main", type=Path, required=True)
    parser.add_argument("--supplement", type=Path, required=True)
    parser.add_argument("--main-pdf", type=Path, required=True)
    parser.add_argument("--supplement-pdf", type=Path, required=True)
    parser.add_argument("--figure-dir", type=Path, required=True)
    parser.add_argument("--basin-csv", type=Path, required=True)
    parser.add_argument("--report", type=Path)
    args = parser.parse_args()
    result = validate(args)
    rendered = json.dumps(result, indent=2, ensure_ascii=False)
    if args.report is not None:
        args.report.parent.mkdir(parents=True, exist_ok=True)
        args.report.write_text(rendered, encoding="utf-8")
    print(rendered)
