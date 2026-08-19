from __future__ import annotations

import argparse
import csv
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


TITLE = "Supplementary Information"
MANUSCRIPT_TITLE = (
    "Climate and Aggregate Human-Activity Scenario Contrasts in Flood and Drought "
    "Event Frequencies across China (1950–2019): An ISIMIP3a Multi-Model Analysis"
)
BASE = Path(__file__).resolve().parent
OUT = BASE / "outputs"
COMMON = OUT / "common_reference_aggregated"
COMPARE = OUT / "calibration_comparison"
VALIDATION = OUT / "validation" / "common_reference" / "model_output_validation.csv"
FIGURE_S1 = OUT / "figures" / "common_reference" / "FigureS_DroughtSensitivity.png"


def configure_analysis_dir(analysis_dir: Path) -> None:
    global COMMON, COMPARE, VALIDATION, FIGURE_S1
    COMMON = analysis_dir / "common_reference_aggregated"
    COMPARE = analysis_dir / "calibration_comparison"
    VALIDATION = analysis_dir / "validation" / "common_reference" / "model_output_validation.csv"
    FIGURE_S1 = analysis_dir / "figures" / "common_reference" / "FigureS_DroughtSensitivity.png"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def f3(value: str | float) -> str:
    return f"{float(value):.3f}"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def format_run(run, size: float = 10.5, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Arial"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Arial")
    rpr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(5)
    format_run(p.add_run(text))


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(11 if level == 1 else 7)
    p.paragraph_format.space_after = Pt(5)
    format_run(p.add_run(text), size=11 if level == 1 else 10.5, bold=True)


def caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    format_run(p.add_run(text), size=9)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, "D9EAF2")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_run(p.add_run(text), size=8, bold=True)
        if widths:
            cell.width = Cm(widths[i])
    for row_values in rows:
        row = table.add_row()
        for i, text in enumerate(row_values):
            cell = row.cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            format_run(p.add_run(str(text)), size=8)
            if widths:
                cell.width = Cm(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def primary_table_rows() -> list[list[str]]:
    rows = read_csv(COMMON / "table1_basin_summary.csv")
    order = {"CSL": 0, "PSE": 1, "SWRB": 2, "YHH": 3, "YZR": 4}
    return [
        [
            r["Basin_Name"], r["Event_Type"],
            f3(r["Mean_Delta_CC_mean"]), f3(r["Mean_Delta_CC_std"]),
            f3(r["Mean_Delta_HA_mean"]), f3(r["Mean_Delta_HA_std"]),
            f3(r["Pooled_DIndex"]),
            f"{r['Climate_Dominant_Model_Count']}/{r['Human_Dominant_Model_Count']}",
        ]
        for r in sorted(rows, key=lambda x: (order[x["Basin_Name"]], x["Event_Type"]))
    ]


def qc_table_rows() -> tuple[list[list[str]], Counter[int]]:
    rows = [r for r in read_csv(VALIDATION) if r["Scenario"] == "obsclim-histsoc"]
    model_rows = [
        [
            r["Model"], r["Calendar"], r["Failed_Month_Fit_Count"],
            r["Grids_With_Failed_Months"], r["Valid_Grid_Count"],
        ]
        for r in rows
    ]
    ensemble = [
        r for r in read_csv(COMMON / "grid_ensemble_attribution.csv")
        if r["Event_Type"] == "Drought"
    ]
    counts = Counter(int(float(r["Delta_Total_count"])) for r in ensemble)
    return model_rows, counts


def robustness_table_rows() -> list[list[str]]:
    primary = {
        (r["Basin_Name"], r["Event_Type"]): float(r["Pooled_DIndex"])
        for r in read_csv(COMMON / "table1_basin_summary.csv")
    }
    defs = read_csv(COMMON / "definition_robustness_range.csv")
    ranges = {
        (r["Family"], r["Basin_Name"], r["Event_Type"]): (float(r["DIndex_Min"]), float(r["DIndex_Max"]))
        for r in defs
    }
    lomo: dict[tuple[str, str], list[float]] = defaultdict(list)
    for r in read_csv(COMMON / "lomo_basin_dindex.csv"):
        lomo[(r["Basin_Name"], r["Event_Type"])].append(float(r["Pooled_DIndex"]))
    order = {"CSL": 0, "PSE": 1, "SWRB": 2, "YHH": 3, "YZR": 4}
    out_rows = []
    for basin, event in sorted(primary, key=lambda x: (order[x[0]], x[1])):
        event_range = ranges[("Event_Definition", basin, event)]
        period_range = ranges[("Period_Split", basin, event)]
        lomo_values = lomo[(basin, event)]
        out_rows.append([
            basin, event, f"{primary[(basin, event)]:.3f}",
            f"{event_range[0]:.3f}–{event_range[1]:.3f}",
            f"{period_range[0]:.3f}–{period_range[1]:.3f}",
            f"{min(lomo_values):.3f}–{max(lomo_values):.3f}",
        ])
    return out_rows


def calibration_basin_rows() -> list[list[str]]:
    rows = read_csv(COMPARE / "calibration_basin_comparison.csv")
    order = {"CSL": 0, "PSE": 1, "SWRB": 2, "YHH": 3, "YZR": 4}
    return [
        [
            r["Basin_Name"], r["Event_Type"],
            f3(r["Pooled_DIndex_ScenarioSpecific"]),
            f3(r["Pooled_DIndex_CommonReference"]),
            f3(r["DIndex_Difference"]),
            "Yes" if r["Dominance_Stable"].lower() == "true" else "No",
        ]
        for r in sorted(rows, key=lambda x: (order[x["Basin_Name"]], x["Event_Type"]))
    ]


def calibration_grid_rows() -> list[list[str]]:
    rows = read_csv(COMPARE / "calibration_grid_ensemble_metrics.csv")
    labels = {"Delta_Total_mean": "ΔTotal", "Delta_CC_mean": "ΔCC", "Delta_HA_mean": "ΔHA"}
    order = {"Delta_Total_mean": 0, "Delta_CC_mean": 1, "Delta_HA_mean": 2}
    return [
        [
            r["Event_Type"], labels[r["Metric"]], r["Paired_Count"],
            f3(r["Pearson_R"]), f3(r["Median_Absolute_Difference"]),
            f3(r["Sign_Concordance"]),
        ]
        for r in sorted(rows, key=lambda x: (x["Event_Type"], order[x["Metric"]]))
    ]


def build(output: Path, analysis_dir: Path = OUT) -> None:
    configure_analysis_dir(analysis_dir)
    required = [
        COMMON / "table1_basin_summary.csv",
        COMMON / "definition_robustness_range.csv",
        COMMON / "lomo_basin_dindex.csv",
        COMMON / "grid_ensemble_attribution.csv",
        COMPARE / "calibration_basin_comparison.csv",
        COMPARE / "calibration_grid_ensemble_metrics.csv",
        VALIDATION,
        FIGURE_S1,
    ]
    missing = [str(p) for p in required if not p.exists()]
    if missing:
        raise FileNotFoundError("Missing required analysis outputs: " + "; ".join(missing))

    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    sec.left_margin = sec.right_margin = Cm(1.8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    format_run(p.add_run(TITLE), size=15, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    format_run(p.add_run(MANUSCRIPT_TITLE), size=11, italic=True)

    body_paragraph(doc, "This file documents the reference-calibrated SRI workflow, quality control, and sensitivity analyses supporting the main manuscript. All numerical tables are generated directly from the final analysis CSV files; ranges are robustness envelopes rather than confidence intervals.")

    heading(doc, "Supplementary Methods S1. Reference-calibrated SRI and event detection")
    body_paragraph(doc, "For each hydrological model and mainland-China grid cell, a 30-day backward-looking running mean was calculated from daily runoff. The factual obsclim-histsoc series supplied the common monthly calibration reference for all three scenarios. For each calendar month over 1950–2019, lognormal, gamma, Gumbel, and Weibull candidates were fitted to positive factual values after division by the positive monthly median; the finite fit with the smallest Akaike information criterion was selected. The fitted family, parameters, zero-flow probability, and numerical scale were then applied without refitting to obsclim-histsoc, counterclim-histsoc, and counterclim-1901soc. This preserves a common runoff-to-SRI mapping within each model-grid and prevents scenario-specific standardization from redefining the event scale.")
    body_paragraph(doc, "Zero runoff was represented by an empirical probability mass and positive runoff by the selected continuous distribution. Mixed cumulative probabilities were converted to standard-normal scores. A model-grid was excluded if any monthly reference fit failed or a non-terminal SRI value remained missing. The domain-wide missing terminal value in JULES-W2 on 2019-12-31 was retained without invalidating an otherwise complete series. No interpolation, random filling, capping, or winsorization was used.")
    body_paragraph(doc, "Events were identified separately within each model and scenario; daily SRI was never averaged across models before event detection. The primary drought definition was SRI ≤ -0.5 for at least 20 consecutive days. The primary flood definition was SRI ≥ 0.5, with pulses separated by fewer than 20 consecutive non-flood days merged into one event. Frequency ratios compared 1950–1984 with 1985–2019 using FR = (N2 + 0.5)/(N1 + 0.5).")

    heading(doc, "Supplementary Methods S2. Scenario contrasts and aggregation")
    body_paragraph(doc, "Within each model-grid, ΔHA = FRcounterclim-histsoc − FRcounterclim-1901soc, ΔCC = FRobsclim-histsoc − FRcounterclim-histsoc, and ΔTotal = FRobsclim-histsoc − FRcounterclim-1901soc. Thus ΔTotal = ΔCC + ΔHA by construction. The terms are scenario contrasts: ΔHA combines model-represented socio-economic differences and does not isolate reservoirs, irrigation, land-use change, groundwater use, or abstraction; ΔCC may include interactions between changing climate and historically modified land-water systems.")
    body_paragraph(doc, "Paired components were first calculated within models and were then summarized with equal weights across available models. Basin means used cosine-latitude weights. Pooled DIndex was calculated from the ensemble means of model-level basin absolute magnitudes as |ΔCC|/(|ΔCC| + |ΔHA|). It is a relative-magnitude diagnostic rather than a percentage contribution.")

    heading(doc, "Supplementary Methods S3. Reservoir-intensity preprocessing")
    body_paragraph(doc, "Reservoir intensity was calculated from the Global Dam Watch v1.0 reservoir data as total mapped storage capacity divided by basin area and expressed as 10⁶ m³ km⁻². Capacity values were coerced to numeric, records without usable capacity were excluded, and only reservoirs spatially joined to one of the five basin groups were retained. Reservoir intensity is a descriptive capacity metric and was not entered into the attribution decomposition or interpreted as causal evidence.")

    heading(doc, "Supplementary Results S1. Quality control and model availability")
    qc_rows, model_counts = qc_table_rows()
    caption(doc, "Table S1. Common-reference calibration quality control by hydrological model. Failed fits and valid-cell counts are identical across the three scenarios because all transformations use the factual monthly reference.")
    add_table(doc, ["Model", "Calendar", "Failed month fits", "Grids with failed fits", "Valid grid cells"], qc_rows, [3.5, 3.2, 3.0, 3.4, 2.8])
    expected_counts = {3: 18, 4: 90, 5: 685, 6: 1287, 7: 1743}
    if dict(model_counts) != expected_counts:
        raise RuntimeError(f"Unexpected valid-model count distribution: {dict(model_counts)}")
    body_paragraph(doc, "All 3,823 coordinate-aligned candidate cells were retained as quality-control rows. After reference-fit exclusions, the number of valid cells ranged from 2,103 to 3,823 among models. At the attribution-ensemble stage, 18, 90, 685, 1,287, and 1,743 cells contained 3, 4, 5, 6, and 7 paired models, respectively. These counts sum to 3,823 and are identical for flood and drought outputs.")

    heading(doc, "Supplementary Results S2. Primary basin estimates")
    caption(doc, "Table S2. Signed common-reference basin components and pooled DIndex. CC/HA model counts show how many of the seven models had the larger absolute climate or aggregate human-activity component.")
    add_table(doc, ["Basin", "Event", "ΔCC mean", "ΔCC SD", "ΔHA mean", "ΔHA SD", "Pooled DIndex", "CC/HA models"], primary_table_rows(), [1.8, 2.0, 2.0, 1.8, 2.0, 1.8, 2.2, 2.2])

    heading(doc, "Supplementary Results S3. Definition, period, and model sensitivity")
    body_paragraph(doc, "Event-definition sensitivity used a 3 × 3 factorial grid for each hazard. Drought thresholds were -0.5, -0.8, and -1.0 with minimum durations of 10, 20, and 30 days. Flood thresholds were 0.5, 0.8, and 1.0 with reset gaps of 10, 20, and 30 consecutive non-flood days. Period sensitivity used split years 1980, 1985, and 1990 and annual event rates for unequal windows. Leave-one-model-out analysis repeated the basin aggregation seven times per basin-event combination.")
    caption(doc, "Table S3. Pooled DIndex under the primary definition and three robustness families. Each range is the minimum–maximum across the tested definitions, period splits, or omitted models.")
    add_table(doc, ["Basin", "Event", "Primary", "9 event definitions", "3 period splits", "7 leave-one-model-out runs"], robustness_table_rows(), [1.7, 1.8, 1.7, 3.2, 3.0, 3.8])
    body_paragraph(doc, "All primary basin classifications remained on the same side of 0.5 across period splits and leave-one-model-out runs. Across the nine event definitions, every classification was stable except YHH flood, which crossed 0.5 (0.459–0.555) and is therefore definition-sensitive despite a primary pooled value of 0.547. YHH drought remained close to parity (0.459–0.478 across event definitions; 0.453–0.496 across period splits; 0.467–0.486 in leave-one-model-out analysis).")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIGURE_S1), width=Cm(16.8))
    caption(doc, "Figure S1. Drought event-definition sensitivity under the common factual SRI reference. Values summarize the share of valid grid cells with FR > 1 across models for the nine threshold-duration combinations. The displayed ranges describe methodological sensitivity and are not confidence intervals.")

    heading(doc, "Supplementary Results S4. Calibration-scheme sensitivity")
    body_paragraph(doc, "The primary analysis uses the common factual SRI reference. A secondary workflow instead fitted each model-scenario-grid series separately. Factual obsclim-histsoc event outputs were exactly identical between workflows (seven files; maximum numeric difference 0; no mismatch in missingness or fitted-distribution labels). Basin dominance classifications were stable for all ten basin-event combinations, but calibration choice affected the fine-scale maps, particularly for floods.")
    caption(doc, "Table S4. Basin pooled DIndex under scenario-specific and common factual calibration. Difference = common reference − scenario-specific.")
    add_table(doc, ["Basin", "Event", "Scenario-specific", "Common reference", "Difference", "Dominance stable"], calibration_basin_rows(), [2.0, 2.0, 3.0, 3.0, 2.2, 2.8])
    doc.add_page_break()
    caption(doc, "Table S5. Grid-ensemble agreement between calibration schemes. Pearson r and sign concordance compare the 3,823 paired grid-cell ensemble means; median absolute difference is in frequency-ratio contrast units.")
    add_table(doc, ["Event", "Component", "Paired cells", "Pearson r", "Median absolute difference", "Sign concordance"], calibration_grid_rows(), [2.0, 2.2, 2.3, 2.2, 3.5, 2.8])
    body_paragraph(doc, "Drought ΔTotal was comparatively stable between calibration schemes (r = 0.882; sign concordance = 0.802). Flood grid-scale patterns were not stable: r was -0.172 for ΔTotal and -0.225 for ΔCC, with sign concordance of 0.599 and 0.582, respectively. Therefore, the manuscript treats common-reference maps as the primary descriptive representation but does not use fine-scale map correspondence as evidence for a robust local mechanism.")

    heading(doc, "Supplementary Methods S4. Reproducibility inventory")
    body_paragraph(doc, "The automated workflow stores model-level event tables, validation tables, per-model attribution, basin summaries, grid-ensemble summaries, sensitivity tables, calibration comparisons, and figure source data as CSV or JSON files. The scripts cover monthly calibration and transformation, event detection, scenario pairing, basin aggregation, leave-one-model-out analysis, sensitivity analysis, structural validation, and figure export. Public repository identifiers have not yet been assigned; the reproducibility archive is currently retained by the authors.")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("output", type=Path)
    parser.add_argument("--analysis-dir", type=Path, default=OUT)
    args = parser.parse_args()
    build(args.output, args.analysis_dir)
    print(args.output.resolve())
